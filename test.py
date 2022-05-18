import docx

import tempfile, shutil, os
def create_temporary_copy(path):
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, 'temp_file_name')
    shutil.copy2(path, temp_path)
    return temp_path


mydoc = docx.Document(create_temporary_copy("./a.docx"))
x = mydoc.add_paragraph("WRITTEN CONSENT")

par = ["I, ________________, owner (and/or authorized representative of Florida Limited Liability Company or Florida Corporation owning a LOT in Boca Pointe Community Association) of LOT ________________________________ with legal description of: _______________________________________ in Boca Pointe Community Association Inc. hereby gives consent for revival of the Declaration of Covenants, Conditions, Restrictions and Easements of Boca Pointe Community Association pursuant to section 720.405(6), Florida Statutes.", "Owner (or authorization representative of Florida entity owning a property / LOT in Boca Pointe Community Association",
"Signature: ______________________________",
"Print Name: ____________________      Date:  _____________________",
"Title:Owner",
"Florida Limited Liability Company name (if applicable):  ______________________________"
]
# font style for x
x.runs[0].font.name = 'Garamond'
x.runs[0].font.size = docx.shared.Pt(12)
x.runs[0].font.bold = True
# align center
x.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

for i in par:
    # font style for x
    p = mydoc.add_paragraph(i)
    p.runs[0].font.name = 'Garamond'
    p.runs[0].font.size = docx.shared.Pt(12)
    # justify
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # align center
# mydoc.save("./b.docx")


from docx2pdf import convert

convert("b.docx")
convert("b.docx", "b.pdf")
