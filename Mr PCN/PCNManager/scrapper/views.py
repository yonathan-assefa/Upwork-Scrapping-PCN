
from django.shortcuts import redirect, render

from django.views.generic import View
from django.views.generic.list import ListView
from urllib3 import HTTPResponse


import docx

import tempfile, shutil, os

from selenium.webdriver.common.by import By

from bs4 import BeautifulSoup
import time
import re
import requests

import pandas as pd

# load variables from __init__.py
from scrapper import browser

from .forms import PCNForm, EmailForm
from .models import PCN, directory

from django.conf import settings

from docx2pdf import convert 
from selenium.common.exceptions import InvalidSessionIdException
def home(request):
    return render(request, "base.html")

def scrap_data(pcn_list):
    pcn_numbers = pcn_list
    for pcn_number in pcn_numbers:
        print("Fetching data for " + pcn_number)
        browser.get(f'https://www.pbcgov.org/papa/Asps/GeneralAdvSrch/NewSearchResults.aspx?srchType=MASTER&proptype=RE&srchVal={pcn_number}&srchPCN=')
        print("Success getting the browser")
        loops = (int(browser.find_element(By.XPATH, '/html/body/form/div[3]/div/div/div[1]/table/tbody/tr[1]/td/table/tbody/tr/td[2]/i/b[3]').text) // 100) + 1 
        for pages in range(loops):
            try:
                print("Searching...")
                s100 = browser.find_element(By.NAME, 'ctl00$MainContent$ddlPgSizeTop').click()
                s100 = browser.find_element(By.NAME, 'ctl00$MainContent$ddlPgSizeTop').send_keys('100')
                print("Sorted by 100")
                time.sleep(2)

                soup = BeautifulSoup(browser.page_source, 'html.parser')
                time.sleep(2)

                sc_list = []

                for row in soup.find_all('tr', class_='gridrow2'):
                    t_datas = row.find_all('td')
                    sc_list.append(t_datas[3].text)

                time.sleep(2)

                for row in soup.find_all('tr', class_='gridrowalternate2'):
                    t_datas = row.find_all('td')
                    sc_list.append(t_datas[3].text)
                print(sc_list)

                def get_user_info(pc_num):
                    each_url = f'https://www.pbcgov.org/papa/Asps/PropertyDetail/PropertyDetail.aspx?parcel={pc_num}&srchtype=MASTER&srchVal=00-42-47-27-38'
                    try:
                        response = requests.get(each_url)
                        soup = BeautifulSoup(response.text, 'html.parser')
                        # find specific id
                        location_address = soup.find('span', id='MainContent_lblLocation').text
                        pcn = soup.find('span', id='MainContent_lblPCN').text
                        subdivision = soup.find('span', id='MainContent_lblSubdiv').text
                        legal_description = soup.find('span', id='MainContent_lblLegalDesc').text

                        sc_list = []
            
                        tbs = soup.find_all("table", {'width':'100%', 'cellspacing':'0', 'cellpadding':'1'})
                        print(tbs)
                    except:
                        print("page unreachable")
                    for i in tbs:
                        for m in i.find_all('tr'):
                            for j in m.find_all('td', {'class': 'TDValueLeft', 'colspan': '2'}):
                                sc_list.append(j.text)

                    owners_name = []
                    for i in sc_list:
                        if i.startswith('\n'):
                            break
                        else:
                            owners_name.append(i)
                    for owner_name in owners_name:
                        # save the data to the PCN table
                        if '&' in owner_name:
                            owner_name = owner_name.replace('&', '')
                        print(owner_name)
                        try:
                            # if the data exists in the database, overwrite it
                            data = PCN.objects.filter(pcn=pcn, owner=owner_name, location=location_address, control_number=pc_num)
                            if data:
                                data.update(subdivision=subdivision, legal_description=legal_description)
                                print("already exists")
                            else:
                                print("saving new")
                                to_be = PCN.objects.create(
                                    pcn = pcn,
                                    owner = owner_name,
                                    location = location_address,
                                    subdivision = subdivision,
                                    legal_description = legal_description,
                                    control_number = pc_num
                                )
                                
                                # for i in owner_name.split():
                                e_p_data = directory.objects.filter(address__icontains=location_address.split()[0])[0]
                                print("found ",  e_p_data)
                                if e_p_data:
                                    to_be.email = e_p_data.email
                                    to_be.phone = f"{e_p_data.phone1}, {e_p_data.phone2}"
                
                                to_be.save()
                        except Exception as e:
                            print(e, "error")
                            pass
                        
                for pc_nums in sc_list:
                    time.sleep(2)
                    get_user_info(pc_nums)

                browser.find_element(By.PARTIAL_LINK_TEXT,value="Next").click()
            except InvalidSessionIdException:
                self.driver.close() 
                print(">>>>> Selenuim just closed the winwdows..,")
            # except:
            #     print("Something went wrong")



class ScrapView(View):
    template_name = 'index.html'
    success_url = 'home'

    def get(self, request):
        form = PCNForm(request.POST or None, request.FILES or None)
        return render(request, self.template_name, {'form': form})
    
    def post(self, request):
        form = PCNForm(request.POST, request.FILES)
        if form.is_valid():
            form = form.cleaned_data
            pcn = form.get('pcn')
            pcn_file = form.get('pcn_file')
            pcn_list = []
            if pcn_file:
                pcn_csv = pd.read_csv(pcn_file)
                pcn_list = pcn_csv['pcn'].tolist()
            if pcn:
                pcn_list.append(pcn)        
            
            scrap_data(pcn_list)
        form = PCNForm(request.POST, request.FILES)
        return render(request, self.template_name, {'form': form})

class FillEmailView(View):
    template_name = 'fill_email.html'
    success_url = 'home'

    def get(self, request):
        form = EmailForm(request.POST or None, request.FILES or None)
        return render(request, self.template_name, {'form': form})
    
    def post(self, request):
        form = EmailForm(request.POST, request.FILES)
        if form.is_valid():
            form = form.cleaned_data
            email_file = request.FILES['email']
            email_csv = pd.read_csv(email_file)
            email_list = email_csv['email'].tolist()
            owner_list = email_csv['owner'].tolist()
            for owner, email in zip(owner_list, email_list):
                PCN.objects.filter(owner=owner).update(email=email)
        form = EmailForm(request.POST, request.FILES)
        return render(request, self.template_name, {'form': form})

class PCNListView(ListView):
    model = PCN
    paginate_by = 100
    template_name = 'pcn_list.html'
    context_object_name = 'pcn_list'


def generate_pdf(request):
# convert html file to pdf using weasyprint
    def create_temporary_copy(path):
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, 'temp_file_name')
        shutil.copy2(path, temp_path)
        return temp_path

    data = PCN.objects.all()
    FLNAME = settings.FLNAME
    for i in data:
        mydoc = docx.Document(create_temporary_copy(FLNAME))
        h = mydoc.add_paragraph("WRITTEN CONSENT")
        # bold h
        h.runs[0].bold = True
        h.runs[0].font.name = 'Garamond'
        h.runs[0].font.size = docx.shared.Pt(12)
        h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        x =  mydoc.add_paragraph("")
        f_name = ' '.join([i.owner.split()[1::-1][0]] + i.owner.split()[2:] +  [i.owner.split()[1::-1][1]])
        print(">>"*5,f_name)
        par = ["I, ", f_name," owner (and/or authorized representative of Florida Limited Liability Company or Florida Corporation owning a LOT in Boca Pointe Community Association) of LOT ", str(i.location), " with legal description of: ",' '.join(i.legal_description.split()), " in Boca Pointe Community Association Inc. hereby gives consent for revival of the Declaration of Covenants, Conditions, Restrictions and Easements of Boca Pointe Community Association pursuant to section 720.405(6), Florida Statutes.", "Owner (or authorization representative of Florida entity owning a property / LOT in Boca Pointe Community Association",
        "Signature: ______________________________",
        "Print Name: ",f_name,"      Date:  _____________________",
        "Title:Owner",
        "Florida Limited Liability Company name (if applicable):  ______________________________"
        ]
        # font style for x

        # align center

        for j in par:
            # font style for x
            # avoid new lines
            if j.startswith("Owner") or j.startswith("Signature") or j.startswith("Florida Limited Liability") or j.startswith("Print Name:") or j.startswith("Title:Owner"):
                x =  mydoc.add_paragraph("")
            try:
                if j.startswith(f_name) or j.startswith(i.location.split()[0]) or j.startswith(i.legal_description.split()[0]):
                    run = x.add_run(j)
                    # underline run
                    run.font.underline = True
                    # p.runs[0].font.bold = True
                    # p.runs[0].font.underline = True
                    run.font.name = 'Garamond'
                    run.font.size = docx.shared.Pt(12)
                    # justify
                    # run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    run = x.add_run(j)
                    run.font.name = 'Garamond'
                    run.font.size = docx.shared.Pt(12)
                    # justify
                    # run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
                # justify x
                x.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            except:
                pass
            #bold
            # align center
        # save mydoc as pdf

        try:
            mydoc.save(f"{settings.PDF_PATH}/{f_name}_{i.location}.docx")
            # convert docx to pdf
            # doc = docx.Document(f"{i.owner}_{i.location}.docx")
            convert(f"{settings.PDF_PATH}/{f_name}_{i.location}.docx", f"{settings.PDF_PATH}/{'_'.join(f_name.split())}.pdf")
            # doc.save(f"{settings.PDF_PATH}/{i.owner}_{i.location}.pdf")
            # delete docx file
            try:
                os.remove(f"{settings.PDF_PATH}/{f_name}_{i.location}.docx")
            except:
                pass
            # send pdf to user
            # mydoc.save(f"{settings.PDF_PATH}/{'_'.join(i.owner.split())}.pdf", format='application/pdf')

        except:
            pass        

        #convert a single docx file to pdf file in same directory

        print('>>>>>>>>>>>>>>>>>>>')
        #convert docx to pdf specifying input & output paths
        # convert(FLNAME,f'{settings.PDF_PATH}/{i.owner}.pdf')
        print('>>>>>>>>>>>>>>>>>>')
    return HTTPResponse("...Already Done!")
      